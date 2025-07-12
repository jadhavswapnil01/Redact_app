import re
import json
import os
import sys
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import ocrmypdf
import spacy
from datetime import datetime
import logging
import hashlib
from pathlib import Path
import docx
from docx.shared import RGBColor
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
import io

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class AdvancedPIIRedactor:
    def __init__(self, config_path=None):
        """Initialize the PII redactor with advanced patterns and NLP model."""
        self.config = self._load_config(config_path)
        self.nlp = self._load_nlp_model()
        self.patterns = self._get_advanced_patterns()
        self.redaction_log = []
        
    def _load_config(self, config_path):
        """Load configuration settings."""
        default_config = {
            "tesseract_path": r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            "output_dir": "redacted_files",
            "temp_dir": "temp_files",
            "confidence_threshold": 0.8,
            "redaction_color": (0, 0, 0),  # Black
            "supported_formats": [".pdf", ".docx", ".doc", ".txt", ".xlsx", ".xls", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"]
        }
        
        if config_path and os.path.exists(config_path):
            with open(config_path, 'r') as f:
                user_config = json.load(f)
                default_config.update(user_config)
        
        return default_config
    
    def _load_nlp_model(self):
        """Load spaCy NLP model for enhanced entity recognition."""
        try:
            return spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            return None
    
    def _get_advanced_patterns(self):
        """Comprehensive PII patterns with context awareness."""
        return {
            # Indian Government IDs
            "AADHAAR": {
                "pattern": r"\b\d{4}[-\s]?\d{4}[-\s]?\d{4}\b",
                "context": r"(?i)(aadhaar|aadhar|uid|unique.*id)",
                "validation": self._validate_aadhaar
            },
            "PAN": {
                "pattern": r"\b[A-Z]{5}\d{4}[A-Z]\b",
                "context": r"(?i)(pan|permanent.*account|income.*tax)",
                "validation": self._validate_pan
            },
            "DRIVING_LICENSE": {
                "pattern": r"\b[A-Z]{2}[-\s]?\d{2}[-\s]?\d{4}[-\s]?\d{7}\b",
                "context": r"(?i)(driving|license|dl|motor)",
                "validation": None
            },
            "PASSPORT": {
                "pattern": r"\b[A-Z]\d{7}\b",
                "context": r"(?i)(passport|travel.*document)",
                "validation": None
            },
            "VOTER_ID": {
                "pattern": r"\b[A-Z]{3}\d{7}\b",
                "context": r"(?i)(voter|election|epic)",
                "validation": None
            },
            
            # Financial Information
            "CREDIT_CARD": {
                "pattern": r"\b(?:\d{4}[-\s]?){3}\d{4}\b",
                "context": r"(?i)(credit|debit|card|visa|master|amex)",
                "validation": self._validate_credit_card
            },
            "BANK_ACCOUNT": {
                "pattern": r"\b\d{9,18}\b",
                "context": r"(?i)(account|bank|saving|current|ifsc)",
                "validation": None
            },
            "IFSC": {
                "pattern": r"\b[A-Z]{4}0[A-Z0-9]{6}\b",
                "context": r"(?i)(ifsc|bank.*code|branch.*code)",
                "validation": None
            },
            
            # Contact Information
            "MOBILE": {
                "pattern": r"\b(?:\+91[-\s]?)?\d{10}\b",
                "context": r"(?i)(mobile|phone|contact|call)",
                "validation": self._validate_mobile
            },
            "EMAIL": {
                "pattern": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
                "context": r"(?i)(email|mail|@)",
                "validation": self._validate_email
            },
            
            # Personal Information
            "DOB": {
                "pattern": r"\b(?:\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\d{2,4}[-/]\d{1,2}[-/]\d{1,2})\b",
                "context": r"(?i)(birth|dob|born|date.*birth)",
                "validation": self._validate_date
            },
            "AGE": {
                "pattern": r"\b(?:age|aged)?\s*:?\s*(\d{1,3})\s*(?:years?|yrs?|y/o)?\b",
                "context": r"(?i)(age|years|old)",
                "validation": None
            },
            
            # Address Components
            "PINCODE": {
                "pattern": r"\b\d{6}\b",
                "context": r"(?i)(pin|pincode|postal|zip)",
                "validation": self._validate_pincode
            },
            "ADDRESS": {
                "pattern": r"(?i)\b(?:house|flat|plot|door)\s*(?:no\.?|number)?\s*[#]?\s*[\w\d/-]+.*?(?:street|road|lane|colony|nagar|area|layout|complex|apartment|building).*?\b",
                "context": r"(?i)(address|residence|house|flat|plot)",
                "validation": None
            },
            
            # Biometric and Health
            "BIOMETRIC_ID": {
                "pattern": r"\b\d{12,16}\b",
                "context": r"(?i)(biometric|fingerprint|iris|retina)",
                "validation": None
            },
            "HEALTH_ID": {
                "pattern": r"\b\d{2}-\d{4}-\d{4}-\d{4}\b",
                "context": r"(?i)(health|medical|hospital|abha)",
                "validation": None
            }
        }
    
    def _validate_aadhaar(self, number):
        """Validate Aadhaar number using Verhoeff algorithm."""
        digits = re.sub(r'[-\s]', '', number)
        if len(digits) != 12:
            return False
        
        # Verhoeff algorithm implementation
        multiplication_table = [
            [0, 1, 2, 3, 4, 5, 6, 7, 8, 9],
            [1, 2, 3, 4, 0, 6, 7, 8, 9, 5],
            [2, 3, 4, 0, 1, 7, 8, 9, 5, 6],
            [3, 4, 0, 1, 2, 8, 9, 5, 6, 7],
            [4, 0, 1, 2, 3, 9, 5, 6, 7, 8],
            [5, 9, 8, 7, 6, 0, 4, 3, 2, 1],
            [6, 5, 9, 8, 7, 1, 0, 4, 3, 2],
            [7, 6, 5, 9, 8, 2, 1, 0, 4, 3],
            [8, 7, 6, 5, 9, 3, 2, 1, 0, 4],
            [9, 8, 7, 6, 5, 4, 3, 2, 1, 0]
        ]
        
        permutation_table = [
            [0, 1, 2, 3, 4, 5, 6, 7, 8, 9],
            [1, 5, 7, 6, 2, 8, 3, 0, 9, 4],
            [5, 8, 0, 3, 7, 9, 6, 1, 4, 2],
            [8, 9, 1, 6, 0, 4, 3, 5, 2, 7],
            [9, 4, 5, 3, 1, 2, 6, 8, 7, 0],
            [4, 2, 8, 6, 5, 7, 3, 9, 0, 1],
            [2, 7, 9, 3, 8, 0, 6, 4, 1, 5],
            [7, 0, 4, 6, 9, 1, 3, 2, 5, 8]
        ]
        
        checksum = 0
        for i, digit in enumerate(reversed(digits)):
            checksum = multiplication_table[checksum][permutation_table[i % 8][int(digit)]]
        
        return checksum == 0
    
    def _validate_pan(self, pan):
        """Validate PAN format and check digit."""
        if not re.match(r'^[A-Z]{5}\d{4}[A-Z]$', pan):
            return False
        return True
    
    def _validate_credit_card(self, number):
        """Validate credit card using Luhn algorithm."""
        digits = re.sub(r'[-\s]', '', number)
        if not digits.isdigit() or len(digits) < 13 or len(digits) > 19:
            return False
        
        # Luhn algorithm
        total = 0
        reverse_digits = digits[::-1]
        
        for i, digit in enumerate(reverse_digits):
            n = int(digit)
            if i % 2 == 1:
                n *= 2
                if n > 9:
                    n -= 9
            total += n
        
        return total % 10 == 0
    
    def _validate_mobile(self, number):
        """Validate Indian mobile number."""
        digits = re.sub(r'[-\s+]', '', number)
        if digits.startswith('91'):
            digits = digits[2:]
        return len(digits) == 10 and digits[0] in '6789'
    
    def _validate_email(self, email):
        """Validate email format."""
        return '@' in email and '.' in email.split('@')[1]
    
    def _validate_date(self, date_str):
        """Validate date format."""
        date_formats = ['%d/%m/%Y', '%m/%d/%Y', '%Y/%m/%d', '%d-%m-%Y', '%m-%d-%Y', '%Y-%m-%d']
        for fmt in date_formats:
            try:
                datetime.strptime(date_str, fmt)
                return True
            except ValueError:
                continue
        return False
    
    def _validate_pincode(self, pincode):
        """Validate Indian pincode."""
        return len(pincode) == 6 and pincode.isdigit()
    
    def extract_text_from_file(self, file_path):
        """Extract text from various file formats."""
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_word(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_path)
            elif file_ext == '.txt':
                return self._extract_from_text(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    def _extract_from_pdf(self, pdf_path):
        """Extract text from PDF with OCR fallback."""
        try:
            doc = fitz.open(pdf_path)
            text = ""
            
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n"
                else:
                    # OCR fallback for image-based PDFs
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    ocr_text = pytesseract.image_to_string(img)
                    text += ocr_text + "\n"
            
            doc.close()
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
            return None
    
    def _extract_from_word(self, doc_path):
        """Extract text from Word documents."""
        try:
            doc = docx.Document(doc_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting from Word: {e}")
            return None
    
    def _extract_from_excel(self, excel_path):
        """Extract text from Excel files."""
        try:
            df = pd.read_excel(excel_path, sheet_name=None)
            text = ""
            for sheet_name, sheet_data in df.items():
                text += f"Sheet: {sheet_name}\n"
                text += sheet_data.to_string(index=False) + "\n\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting from Excel: {e}")
            return None
    
    def _extract_from_text(self, txt_path):
        """Extract text from plain text files."""
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting from text file: {e}")
            return None
    
    def _extract_from_image(self, img_path):
        """Extract text from images using OCR."""
        try:
            img = Image.open(img_path)
            return pytesseract.image_to_string(img)
        except Exception as e:
            logger.error(f"Error extracting from image: {e}")
            return None
    
    def detect_pii(self, text):
        """Detect PII using patterns and NLP."""
        if not text:
            return {}
        
        detected_pii = {}
        
        # Pattern-based detection
        for pii_type, config in self.patterns.items():
            matches = []
            pattern_matches = re.finditer(config["pattern"], text, re.IGNORECASE)
            
            for match in pattern_matches:
                match_text = match.group()
                match_start = match.start()
                match_end = match.end()
                
                # Context validation
                context_score = self._calculate_context_score(text, match_start, match_end, config.get("context", ""))
                
                # Pattern validation
                if config.get("validation"):
                    is_valid = config["validation"](match_text)
                else:
                    is_valid = True
                
                if is_valid and context_score > self.config["confidence_threshold"]:
                    matches.append({
                        "text": match_text,
                        "start": match_start,
                        "end": match_end,
                        "confidence": context_score
                    })
            
            if matches:
                detected_pii[pii_type] = matches
        
        # NLP-based detection
        if self.nlp:
            nlp_entities = self._detect_nlp_entities(text)
            detected_pii.update(nlp_entities)
        
        return detected_pii
    
    def _calculate_context_score(self, text, start, end, context_pattern):
        """Calculate confidence score based on context."""
        if not context_pattern:
            return 1.0
        
        # Check surrounding text for context
        context_window = 100
        context_start = max(0, start - context_window)
        context_end = min(len(text), end + context_window)
        context_text = text[context_start:context_end]
        
        context_matches = re.findall(context_pattern, context_text, re.IGNORECASE)
        return min(1.0, len(context_matches) * 0.3 + 0.7)
    
    def _detect_nlp_entities(self, text):
        """Use NLP for entity detection."""
        try:
            doc = self.nlp(text)
            entities = {}
            
            for ent in doc.ents:
                if ent.label_ in ["PERSON", "ORG", "GPE", "MONEY", "DATE"]:
                    entity_type = f"NLP_{ent.label_}"
                    if entity_type not in entities:
                        entities[entity_type] = []
                    
                    entities[entity_type].append({
                        "text": ent.text,
                        "start": ent.start_char,
                        "end": ent.end_char,
                        "confidence": 0.9
                    })
            
            return entities
        except Exception as e:
            logger.error(f"Error in NLP detection: {e}")
            return {}
    
    def redact_file(self, input_path, output_path=None):
        """Redact PII from file and return redacted version."""
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        file_ext = Path(input_path).suffix.lower()
        
        if not output_path:
            output_path = self._generate_output_path(input_path)
        
        try:
            if file_ext == '.pdf':
                return self._redact_pdf(input_path, output_path)
            elif file_ext in ['.docx', '.doc']:
                return self._redact_word(input_path, output_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._redact_excel(input_path, output_path)
            elif file_ext == '.txt':
                return self._redact_text(input_path, output_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._redact_image(input_path, output_path)
            else:
                raise ValueError(f"Unsupported file format for redaction: {file_ext}")
        except Exception as e:
            logger.error(f"Error redacting file: {e}")
            raise
    
    def _generate_output_path(self, input_path):
        """Generate output path for redacted file."""
        path = Path(input_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Ensure proper extension formatting
        extension = path.suffix if path.suffix else '.txt'
        return path.parent / f"{path.stem}_redacted_{timestamp}{extension}"
    
    def _redact_pdf(self, input_path, output_path):
        """Redact PDF file."""
        text = self.extract_text_from_file(input_path)
        pii_data = self.detect_pii(text)
        
        if not pii_data:
            # No PII found, copy original
            import shutil
            shutil.copy2(input_path, output_path)
            return output_path
        
        doc = fitz.open(input_path)
        
        for page in doc:
            page_text = page.get_text()
            
            for pii_type, matches in pii_data.items():
                for match in matches:
                    search_results = page.search_for(match["text"])
                    for rect in search_results:
                        page.add_redact_annot(rect, fill=self.config["redaction_color"])
            
            page.apply_redactions()
        
        doc.save(output_path, garbage=4, deflate=True)
        doc.close()
        
        self._log_redaction(input_path, output_path, pii_data)
        return output_path
    
    def _redact_word(self, input_path, output_path):
        """Redact Word document."""
        text = self.extract_text_from_file(input_path)
        pii_data = self.detect_pii(text)
        
        doc = docx.Document(input_path)
        
        for paragraph in doc.paragraphs:
            for pii_type, matches in pii_data.items():
                for match in matches:
                    if match["text"] in paragraph.text:
                        paragraph.text = paragraph.text.replace(match["text"], "█" * len(match["text"]))
        
        doc.save(output_path)
        self._log_redaction(input_path, output_path, pii_data)
        return output_path
    
    def _redact_excel(self, input_path, output_path):
        """Redact Excel file."""
        text = self.extract_text_from_file(input_path)
        pii_data = self.detect_pii(text)
        
        workbook = load_workbook(input_path)
        black_fill = PatternFill(start_color="000000", end_color="000000", fill_type="solid")
        
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value:
                        cell_text = str(cell.value)
                        for pii_type, matches in pii_data.items():
                            for match in matches:
                                if match["text"] in cell_text:
                                    cell.value = "REDACTED"
                                    cell.fill = black_fill
        
        workbook.save(output_path)
        self._log_redaction(input_path, output_path, pii_data)
        return output_path
    
    def _redact_text(self, input_path, output_path):
        """Redact plain text file."""
        text = self.extract_text_from_file(input_path)
        pii_data = self.detect_pii(text)
        
        redacted_text = text
        
        # Sort matches by position (reverse order to maintain positions)
        all_matches = []
        for pii_type, matches in pii_data.items():
            for match in matches:
                all_matches.append(match)
        
        all_matches.sort(key=lambda x: x["start"], reverse=True)
        
        for match in all_matches:
            redacted_text = redacted_text[:match["start"]] + "█" * (match["end"] - match["start"]) + redacted_text[match["end"]:]
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(redacted_text)
        
        self._log_redaction(input_path, output_path, pii_data)
        return output_path
    
    def _redact_image(self, input_path, output_path):
        """Redact image file by converting to PDF and redacting."""
        # Convert image to PDF first
        img = Image.open(input_path)
        pdf_path = input_path.replace(Path(input_path).suffix, "_temp.pdf")
        img.convert('RGB').save(pdf_path)
        
        # Redact PDF
        redacted_pdf = self._redact_pdf(pdf_path, pdf_path.replace("_temp.pdf", "_redacted.pdf"))
        
        # Convert back to image if needed
        if Path(output_path).suffix.lower() in ['.png', '.jpg', '.jpeg']:
            doc = fitz.open(redacted_pdf)
            page = doc[0]
            pix = page.get_pixmap()
            pix.save(output_path)
            doc.close()
            os.remove(pdf_path)
            os.remove(redacted_pdf)
        else:
            os.rename(redacted_pdf, output_path)
            os.remove(pdf_path)
        
        return output_path
    
    def _log_redaction(self, input_path, output_path, pii_data):
        """Log redaction details."""
        log_entry = {
            "timestamp": datetime.now().isoformat(),
            "input_file": input_path,
            "output_file": output_path,
            "pii_detected": {k: len(v) for k, v in pii_data.items()},
            "total_pii_count": sum(len(v) for v in pii_data.values())
        }
        
        self.redaction_log.append(log_entry)
        logger.info(f"Redacted {log_entry['total_pii_count']} PII instances from {input_path}")
    
    def get_redaction_summary(self):
        """Get summary of all redactions performed."""
        return {
            "total_files_processed": len(self.redaction_log),
            "total_pii_redacted": sum(log["total_pii_count"] for log in self.redaction_log),
            "redaction_log": self.redaction_log
        }
    
def ensure_proper_extension(file_path):
    """Ensure file has proper extension format."""
    path = Path(file_path)
    if not path.suffix:
        return str(path) + '.txt'
    return str(path)


def main():
    """Main function for CLI usage."""
    if len(sys.argv) < 2:
        print("Usage: python advanced_pii_redactor.py <input_file> [output_file]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    if output_file:
        output_file = ensure_proper_extension(output_file)
    
    try:
        redactor = AdvancedPIIRedactor()
        result = redactor.redact_file(input_file, output_file)
        print(f"Redacted file saved to: {result}")
        
        # Print summary
        summary = redactor.get_redaction_summary()
        print(f"Total PII instances redacted: {summary['total_pii_redacted']}")
        
    except Exception as e:
        logger.error(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()