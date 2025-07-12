import re
import json
import os
import sys
import fitz
import pytesseract
from PIL import Image, ImageDraw, ImageFont
import ocrmypdf
import spacy
from datetime import datetime
import logging
import hashlib
from pathlib import Path
import docx
from docx.shared import RGBColor
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
import io
import cv2
import numpy as np
from collections import defaultdict
import base64
import tempfile

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class AdvancedPIIRedactor:
    def __init__(self, config_path=None):
        """Initialize the PII redactor with comprehensive patterns and NLP model."""
        self.config = self._load_config(config_path)
        self.nlp = self._load_nlp_model()
        self.patterns = self._get_comprehensive_patterns()
        self.redaction_log = []
        self.context_keywords = self._load_context_keywords()
        
    def _load_config(self, config_path):
        """Load configuration settings."""
        default_config = {
            "tesseract_path": r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            "output_dir": "redacted_files",
            "temp_dir": "temp_files",
            "confidence_threshold": 0.6,
            "context_window": 150,
            "redaction_color": (0, 0, 0),
            "supported_formats": [".pdf", ".docx", ".doc", ".txt", ".xlsx", ".xls", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"]
        }
        
        if config_path and os.path.exists(config_path):
            with open(config_path, 'r') as f:
                user_config = json.load(f)
                default_config.update(user_config)
        
        return default_config
    
    def _load_nlp_model(self):
        """Load spaCy NLP model for enhanced entity recognition."""
        try:
            return spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model not found. Using pattern-based detection only.")
            return None
    
    def _load_context_keywords(self):
        """Load context keywords for better PII detection."""
        return {
            "AADHAAR": ["aadhaar", "aadhar", "uid", "unique", "identification", "uidai", "enrollment", "demographic"],
            "PAN": ["pan", "permanent", "account", "number", "income", "tax", "assessee", "taxpayer"],
            "DRIVING_LICENSE": ["driving", "license", "dl", "motor", "vehicle", "transport", "rto", "issued"],
            "PASSPORT": ["passport", "travel", "document", "republic", "india", "issued", "valid"],
            "VOTER_ID": ["voter", "election", "epic", "electoral", "commission", "booth", "constituency"],
            "CREDIT_CARD": ["credit", "debit", "card", "visa", "master", "amex", "discover", "valid", "expires"],
            "BANK_ACCOUNT": ["account", "bank", "saving", "current", "deposit", "branch", "holder"],
            "IFSC": ["ifsc", "micr", "branch", "code", "routing", "swift", "bank"],
            "MOBILE": ["mobile", "phone", "contact", "cell", "number", "registered", "primary"],
            "EMAIL": ["email", "mail", "address", "contact", "registered", "primary", "alternate"],
            "DOB": ["birth", "dob", "born", "date", "age", "year", "month", "day"],
            "ADDRESS": ["address", "residence", "house", "flat", "plot", "street", "road", "city", "state", "country"],
            "PINCODE": ["pin", "pincode", "postal", "zip", "area", "code"],
            "NAME": ["name", "first", "last", "middle", "full", "surname", "given", "family"],
            "FATHER_NAME": ["father", "dad", "papa", "parent", "guardian", "s/o", "son", "daughter"],
            "MOTHER_NAME": ["mother", "mom", "mama", "parent", "guardian", "d/o", "w/o", "wife"],
            "SPOUSE_NAME": ["spouse", "husband", "wife", "partner", "married", "w/o", "s/o"]
        }
    
    def _get_comprehensive_patterns(self):
        """Comprehensive PII patterns with enhanced validation."""
        return {
            # Indian Government IDs
            "AADHAAR": {
                "patterns": [
                    r"\b\d{4}[-\s]?\d{4}[-\s]?\d{4}\b",
                    r"\b\d{12}\b",
                    r"(?i)(?:aadhaar|aadhar|uid)[-\s:]*(?:no|number|#)?[-\s:]*(\d{4}[-\s]?\d{4}[-\s]?\d{4})",
                    r"(?i)(?:aadhaar|aadhar|uid)[-\s:]*(?:no|number|#)?[-\s:]*(\d{12})"
                ],
                "validation": self._validate_aadhaar,
                "priority": 1
            },
            "PAN": {
                "patterns": [
                    r"\b[A-Z]{5}\d{4}[A-Z]\b",
                    r"(?i)(?:pan|permanent.*account)[-\s:]*(?:no|number|#)?[-\s:]*([A-Z]{5}\d{4}[A-Z])"
                ],
                "validation": self._validate_pan,
                "priority": 1
            },
            "DRIVING_LICENSE": {
                "patterns": [
                    r"\b[A-Z]{2}[-\s]?\d{2}[-\s]?\d{4}[-\s]?\d{7}\b",
                    r"\b[A-Z]{2}\d{13}\b",
                    r"(?i)(?:driving.*license|dl)[-\s:]*(?:no|number|#)?[-\s:]*([A-Z]{2}[-\s]?\d{2}[-\s]?\d{4}[-\s]?\d{7})"
                ],
                "validation": self._validate_dl,
                "priority": 2
            },
            "PASSPORT": {
                "patterns": [
                    r"\b[A-Z]\d{7}\b",
                    r"\b[A-Z]{2}\d{6}\b",
                    r"(?i)(?:passport)[-\s:]*(?:no|number|#)?[-\s:]*([A-Z]\d{7})"
                ],
                "validation": self._validate_passport,
                "priority": 2
            },
            "VOTER_ID": {
                "patterns": [
                    r"\b[A-Z]{3}\d{7}\b",
                    r"\b[A-Z]{2}[A-Z0-9]\d{6}\b",
                    r"(?i)(?:voter.*id|epic)[-\s:]*(?:no|number|#)?[-\s:]*([A-Z]{3}\d{7})"
                ],
                "validation": self._validate_voter_id,
                "priority": 2
            },
            
            # Financial Information
            "CREDIT_CARD": {
                "patterns": [
                    r"\b(?:\d{4}[-\s]?){3}\d{4}\b",
                    r"\b\d{13,19}\b",
                    r"(?i)(?:credit|debit|card)[-\s:]*(?:no|number|#)?[-\s:]*(\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4})"
                ],
                "validation": self._validate_credit_card,
                "priority": 1
            },
            "BANK_ACCOUNT": {
                "patterns": [
                    r"\b\d{9,18}\b",
                    r"(?i)(?:account)[-\s:]*(?:no|number|#)?[-\s:]*(\d{9,18})"
                ],
                "validation": self._validate_bank_account,
                "priority": 2
            },
            "IFSC": {
                "patterns": [
                    r"\b[A-Z]{4}0[A-Z0-9]{6}\b",
                    r"(?i)(?:ifsc)[-\s:]*(?:code)?[-\s:]*([A-Z]{4}0[A-Z0-9]{6})"
                ],
                "validation": self._validate_ifsc,
                "priority": 2
            },
            
            # Contact Information
            "MOBILE": {
                "patterns": [
                    r"\b(?:\+91[-\s]?)?\d{10}\b",
                    r"\b91[-\s]?\d{10}\b",
                    r"(?i)(?:mobile|phone|contact)[-\s:]*(?:no|number|#)?[-\s:]*(\+?91[-\s]?\d{10})",
                    r"(?i)(?:mobile|phone|contact)[-\s:]*(?:no|number|#)?[-\s:]*(\d{10})"
                ],
                "validation": self._validate_mobile,
                "priority": 1
            },
            "EMAIL": {
                "patterns": [
                    r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
                    r"(?i)(?:email|mail)[-\s:]*(?:id|address)?[-\s:]*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})"
                ],
                "validation": self._validate_email,
                "priority": 1
            },
            
            # Personal Information
            "DOB": {
                "patterns": [
                    r"\b(?:\d{1,2}[-/]\d{1,2}[-/]\d{2,4})\b",
                    r"\b(?:\d{2,4}[-/]\d{1,2}[-/]\d{1,2})\b",
                    r"(?i)(?:dob|date.*birth|born)[-\s:]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})"
                ],
                "validation": self._validate_date,
                "priority": 1
            },
            "AGE": {
                "patterns": [
                    r"\b(?:age|aged)?\s*:?\s*(\d{1,3})\s*(?:years?|yrs?|y/o)?\b",
                    r"(?i)(?:age)[-\s:]*(\d{1,3})"
                ],
                "validation": self._validate_age,
                "priority": 3
            },
            
            # Address Components
            "PINCODE": {
                "patterns": [
                    r"\b\d{6}\b",
                    r"(?i)(?:pin|pincode|postal)[-\s:]*(?:code)?[-\s:]*(\d{6})"
                ],
                "validation": self._validate_pincode,
                "priority": 2
            },
            "ADDRESS": {
                "patterns": [
                    r"(?i)\b(?:house|flat|plot|door|room|block|building|apartment|complex|colony|nagar|layout|area|sector|phase|street|road|lane|avenue|cross|main|circle|square|park|garden|residency|estate|township|enclave|extension|vihar|puram|city|town|village|dist|district|state|country)[\s\w\d,.-]{5,200}",
                    r"(?i)(?:address|residence)[-\s:]*([^,\n]{10,200})"
                ],
                "validation": self._validate_address,
                "priority": 3
            },
            
            # Names
            "NAME": {
                "patterns": [
                    r"(?i)(?:name|first.*name|last.*name|full.*name)[-\s:]*([A-Za-z\s]{2,50})",
                    r"\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b"
                ],
                "validation": self._validate_name,
                "priority": 3
            },
            "FATHER_NAME": {
                "patterns": [
                    r"(?i)(?:father.*name|s/o|son.*of)[-\s:]*([A-Za-z\s]{2,50})",
                    r"(?i)(?:father)[-\s:]*([A-Za-z\s]{2,50})"
                ],
                "validation": self._validate_name,
                "priority": 2
            },
            "MOTHER_NAME": {
                "patterns": [
                    r"(?i)(?:mother.*name|d/o|daughter.*of|w/o|wife.*of)[-\s:]*([A-Za-z\s]{2,50})",
                    r"(?i)(?:mother)[-\s:]*([A-Za-z\s]{2,50})"
                ],
                "validation": self._validate_name,
                "priority": 2
            },
            
            # Biometric and Health
            "BIOMETRIC_ID": {
                "patterns": [
                    r"\b\d{12,16}\b",
                    r"(?i)(?:biometric|fingerprint|iris)[-\s:]*(?:id|number)?[-\s:]*(\d{12,16})"
                ],
                "validation": self._validate_biometric,
                "priority": 3
            },
            "HEALTH_ID": {
                "patterns": [
                    r"\b\d{2}-\d{4}-\d{4}-\d{4}\b",
                    r"(?i)(?:health.*id|abha|medical.*id)[-\s:]*(\d{2}-\d{4}-\d{4}-\d{4})"
                ],
                "validation": self._validate_health_id,
                "priority": 2
            }
        }
    
    def _validate_aadhaar(self, number):
        """Enhanced Aadhaar validation with Verhoeff algorithm."""
        digits = re.sub(r'[-\s]', '', number)
        if len(digits) != 12 or not digits.isdigit():
            return False
        
        # Check for invalid patterns
        if digits == '0' * 12 or digits == '1' * 12:
            return False
        
        # Verhoeff algorithm
        multiplication_table = [
            [0, 1, 2, 3, 4, 5, 6, 7, 8, 9],
            [1, 2, 3, 4, 0, 6, 7, 8, 9, 5],
            [2, 3, 4, 0, 1, 7, 8, 9, 5, 6],
            [3, 4, 0, 1, 2, 8, 9, 5, 6, 7],
            [4, 0, 1, 2, 3, 9, 5, 6, 7, 8],
            [5, 9, 8, 7, 6, 0, 4, 3, 2, 1],
            [6, 5, 9, 8, 7, 1, 0, 4, 3, 2],
            [7, 6, 5, 9, 8, 2, 1, 0, 4, 3],
            [8, 7, 6, 5, 9, 3, 2, 1, 0, 4],
            [9, 8, 7, 6, 5, 4, 3, 2, 1, 0]
        ]
        
        permutation_table = [
            [0, 1, 2, 3, 4, 5, 6, 7, 8, 9],
            [1, 5, 7, 6, 2, 8, 3, 0, 9, 4],
            [5, 8, 0, 3, 7, 9, 6, 1, 4, 2],
            [8, 9, 1, 6, 0, 4, 3, 5, 2, 7],
            [9, 4, 5, 3, 1, 2, 6, 8, 7, 0],
            [4, 2, 8, 6, 5, 7, 3, 9, 0, 1],
            [2, 7, 9, 3, 8, 0, 6, 4, 1, 5],
            [7, 0, 4, 6, 9, 1, 3, 2, 5, 8]
        ]
        
        checksum = 0
        for i, digit in enumerate(reversed(digits)):
            checksum = multiplication_table[checksum][permutation_table[i % 8][int(digit)]]
        
        return checksum == 0
    
    def _validate_pan(self, pan):
        """Enhanced PAN validation."""
        pan = pan.upper()
        if not re.match(r'^[A-Z]{5}\d{4}[A-Z]$', pan):
            return False
        
        # Check for invalid patterns
        if pan[0:5] == 'AAAAA' or pan[5:9] == '0000':
            return False
        
        return True
    
    def _validate_dl(self, dl):
        """Validate driving license format."""
        dl = re.sub(r'[-\s]', '', dl)
        return len(dl) == 15 and dl[:2].isalpha() and dl[2:].isdigit()
    
    def _validate_passport(self, passport):
        """Validate passport format."""
        return len(passport) >= 7 and passport[0].isalpha()
    
    def _validate_voter_id(self, voter_id):
        """Validate voter ID format."""
        return len(voter_id) == 10 and voter_id[:3].isalpha() and voter_id[3:].isdigit()
    
    def _validate_credit_card(self, number):
        """Enhanced credit card validation using Luhn algorithm."""
        digits = re.sub(r'[-\s]', '', number)
        if not digits.isdigit() or len(digits) < 13 or len(digits) > 19:
            return False
        
        # Luhn algorithm
        total = 0
        reverse_digits = digits[::-1]
        
        for i, digit in enumerate(reverse_digits):
            n = int(digit)
            if i % 2 == 1:
                n *= 2
                if n > 9:
                    n -= 9
            total += n
        
        return total % 10 == 0
    
    def _validate_bank_account(self, account):
        """Validate bank account number."""
        return 9 <= len(account) <= 18 and account.isdigit()
    
    def _validate_ifsc(self, ifsc):
        """Validate IFSC code format."""
        return len(ifsc) == 11 and ifsc[:4].isalpha() and ifsc[4] == '0'
    
    def _validate_mobile(self, number):
        """Enhanced mobile number validation."""
        digits = re.sub(r'[-\s+]', '', number)
        if digits.startswith('91'):
            digits = digits[2:]
        return len(digits) == 10 and digits[0] in '6789' and digits.isdigit()
    
    def _validate_email(self, email):
        """Enhanced email validation."""
        return '@' in email and '.' in email.split('@')[1] and len(email.split('@')) == 2
    
    def _validate_date(self, date_str):
        """Enhanced date validation."""
        date_formats = ['%d/%m/%Y', '%m/%d/%Y', '%Y/%m/%d', '%d-%m-%Y', '%m-%d-%Y', '%Y-%m-%d']
        for fmt in date_formats:
            try:
                datetime.strptime(date_str, fmt)
                return True
            except ValueError:
                continue
        return False
    
    def _validate_age(self, age):
        """Validate age."""
        return 0 <= int(age) <= 150
    
    def _validate_pincode(self, pincode):
        """Enhanced pincode validation."""
        return len(pincode) == 6 and pincode.isdigit() and not pincode.startswith('0')
    
    def _validate_address(self, address):
        """Validate address format."""
        return len(address.strip()) > 10 and any(word in address.lower() for word in ['house', 'flat', 'plot', 'street', 'road', 'city', 'state'])
    
    def _validate_name(self, name):
        """Validate name format."""
        name = name.strip()
        return 2 <= len(name) <= 50 and all(c.isalpha() or c.isspace() for c in name)
    
    def _validate_biometric(self, bio_id):
        """Validate biometric ID."""
        return 12 <= len(bio_id) <= 16 and bio_id.isdigit()
    
    def _validate_health_id(self, health_id):
        """Validate health ID format."""
        return re.match(r'^\d{2}-\d{4}-\d{4}-\d{4}$', health_id) is not None
    
    def extract_text_from_file(self, file_path):
        """Enhanced text extraction with OCR improvements."""
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_word(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_path)
            elif file_ext == '.txt':
                return self._extract_from_text(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    def _extract_from_pdf(self, pdf_path):
        """Enhanced PDF text extraction with better OCR."""
        try:
            doc = fitz.open(pdf_path)
            text = ""
            
            for page in doc:
                # Try direct text extraction first
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n"
                else:
                    # OCR fallback with image preprocessing
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Higher resolution
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Preprocess image for better OCR
                    img = self._preprocess_image_for_ocr(img)
                    ocr_text = pytesseract.image_to_string(img, config='--psm 6')
                    text += ocr_text + "\n"
            
            doc.close()
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
            return None
    
    def _extract_from_image(self, img_path):
        """Enhanced image text extraction with preprocessing."""
        try:
            img = Image.open(img_path)
            img = self._preprocess_image_for_ocr(img)
            return pytesseract.image_to_string(img, config='--psm 6')
        except Exception as e:
            logger.error(f"Error extracting from image: {e}")
            return None
    
    def _preprocess_image_for_ocr(self, img):
        """Preprocess image for better OCR results."""
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Convert to OpenCV format
        img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
        
        # Apply preprocessing
        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        
        # Noise removal
        denoised = cv2.medianBlur(gray, 3)
        
        # Thresholding
        _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Convert back to PIL Image
        return Image.fromarray(thresh)
    
    def _extract_from_word(self, doc_path):
        """Extract text from Word documents."""
        try:
            doc = docx.Document(doc_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting from Word: {e}")
            return None
    
    def _extract_from_excel(self, excel_path):
        """Extract text from Excel files."""
        try:
            df = pd.read_excel(excel_path, sheet_name=None)
            text = ""
            for sheet_name, sheet_data in df.items():
                text += f"Sheet: {sheet_name}\n"
                text += sheet_data.to_string(index=False) + "\n\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting from Excel: {e}")
            return None
    
    def _extract_from_text(self, txt_path):
        """Extract text from plain text files."""
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting from text file: {e}")
            return None
    
    def detect_pii(self, text):
        """Enhanced PII detection with multi-stage approach."""
        if not text:
            return {}
        
        detected_pii = {}
        
        # Stage 1: Pattern-based detection with context
        for pii_type, config in self.patterns.items():
            matches = []
            
            for pattern in config["patterns"]:
                pattern_matches = list(re.finditer(pattern, text, re.IGNORECASE))
                
                for match in pattern_matches:
                    match_text = match.group(1) if match.groups() else match.group()
                    match_start = match.start(1) if match.groups() else match.start()
                    match_end = match.end(1) if match.groups() else match.end()
                    
                    # Context validation
                    context_score = self._calculate_context_score(text, match_start, match_end, pii_type)
                    
                    # Pattern validation
                    if config.get("validation"):
                        try:
                            is_valid = config["validation"](match_text)
                        except:
                            is_valid = False
                    else:
                        is_valid = True
                    
                    # Priority-based confidence adjustment
                    priority_boost = (4 - config.get("priority", 3)) * 0.1
                    final_confidence = min(1.0, context_score + priority_boost)
                    
                    if is_valid and final_confidence > self.config["confidence_threshold"]:
                        matches.append({
                            "text": match_text,
                            "start": match_start,
                            "end": match_end,
                            "confidence": final_confidence,
                            "pattern": pattern
                        })
            
            if matches:
                # Remove duplicates and overlapping matches
                matches = self._remove_overlapping_matches(matches)
                detected_pii[pii_type] = matches
        
        # Stage 2: NLP-based detection
        if self.nlp:
            nlp_entities = self._detect_nlp_entities(text)
            for entity_type, entities in nlp_entities.items():
                if entity_type not in detected_pii:
                    detected_pii[entity_type] = entities
                else:
                    detected_pii[entity_type].extend(entities)
        
        # Stage 3: Post-processing and refinement
        detected_pii = self._refine_detections(detected_pii, text)
        
        return detected_pii
    
    def _calculate_context_score(self, text, start, end, pii_type):
        """Enhanced context scoring with keyword matching."""
        context_window = self.config["context_window"]
        context_start = max(0, start - context_window)
        context_end = min(len(text), end + context_window)
        context_text = text[context_start:context_end].lower()
        
        # Check for relevant keywords
        keywords = self.context_keywords.get(pii_type, [])
        keyword_score = 0
        
        for keyword in keywords:
            if keyword in context_text:
                keyword_score += 0.2
        
        # Base score
        base_score = 0.7
        
        # Combine scores
        final_score = min(1.0, base_score + keyword_score)
        
        return final_score
    
    def _remove_overlapping_matches(self, matches):
        """Remove overlapping matches, keeping the highest confidence ones."""
        if not matches:
            return matches
        
        # Sort by confidence (descending)
        sorted_matches = sorted(matches, key=lambda x: x["confidence"], reverse=True)
        
        filtered_matches = []
        for match in sorted_matches:
            is_overlapping = False
            for existing in filtered_matches:
                if (match["start"] < existing["end"] and match["end"] > existing["start"]):
                    is_overlapping = True
                    break
            
            if not is_overlapping:
                filtered_matches.append(match)
        
        return filtered_matches
    
    def _detect_nlp_entities(self, text):
        """Enhanced NLP entity detection."""
        try:
            doc = self.nlp(text)
            entities = {}
            
            for ent in doc.ents:
                if ent.label_ in ["PERSON", "ORG", "GPE", "MONEY", "DATE", "CARDINAL"]:
                    entity_type = f"NLP_{ent.label_}"
                    if entity_type not in entities:
                        entities[entity_type] = []
                    
                    entities[entity_type].append({
                        "text": ent.text,
                        "start": ent.start_char,
                        "end": ent.end_char,
                        "confidence": 0.8
                    })
            
            return entities
        except Exception as e:
            logger.error(f"Error in NLP detection: {e}")
            return {}
    
    def _refine_detections(self, detected_pii, text):
        """Refine and filter detected PII."""
        refined_pii = {}
        
        for pii_type, matches in detected_pii.items():
            refined_matches = []
            
            for match in matches:
                # Skip very short matches for certain types
                if pii_type in ["NAME", "ADDRESS"] and len(match["text"].strip()) < 3:
                    continue
                
                # Skip common false positives
                if self._is_false_positive(match["text"], pii_type):
                    continue
                
                refined_matches.append(match)
            
            if refined_matches:
                refined_pii[pii_type] = refined_matches
        
        return refined_pii
    
    def _is_false_positive(self, text, pii_type):
        """Check for common false positives."""
        text_lower = text.lower().strip()
        
        # Common false positives
        false_positives = {
            "NAME": ["name", "first", "last", "full", "mr", "mrs", "ms", "dr", "sir", "madam"],
            "MOBILE": ["0000000000", "1111111111", "9999999999"],
            "EMAIL": ["email@example.com", "test@test.com"],
            "ADDRESS": ["address", "street", "city", "state"]
        }
        
        return text_lower in false_positives.get(pii_type, [])
    
    def redact_file(self, input_path, output_path=None):
        """Enhanced file redaction with perfect masking."""
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        file_ext = Path(input_path).suffix.lower()
        
        if not output_path:
            output_path = self._generate_output_path(input_path)
        
        try:
            if file_ext == '.pdf':
                return self._redact_pdf(input_path, output_path)
            elif file_ext in ['.docx', '.doc']:
                return self._redact_word(input_path, output_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._redact_excel(input_path, output_path)
            elif file_ext == '.txt':
                return self._redact_text(input_path, output_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._redact_image(input_path, output_path)
            else:
                raise ValueError(f"Unsupported file format for redaction: {file_ext}")
        except Exception as e:
            logger.error(f"Error redacting file: {e}")
            raise
    
    def _generate_output_path(self, input_path):
        """Generate output path for redacted file."""
        path = Path(input_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        extension = path.suffix if path.suffix else '.txt'
        return path.parent / f"{path.stem}_redacted_{timestamp}{extension}"
    
    def _redact_pdf(self, input_path, output_path):
        """Enhanced PDF redaction with precise positioning."""
        text = self.extract_text_from_file(input_path)
        pii_data = self.detect_pii(text)
        
        if not pii_data:
            import shutil
            shutil.copy2(input_path, output_path)
            return output_path
        
        doc = fitz.open(input_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            
            # Get text blocks with positions
            blocks = page.get_text("dict")
            
            for pii_type, matches in pii_data.items():
                for match in matches:
                    match_text = match["text"]
                    
                    # Find text instances on page
                    text_instances = page.search_for(match_text)
                    
                    for rect in text_instances:
                        # Create redaction annotation
                        redact_annot = page.add_redact_annot(rect)
                        redact_annot.set_colors(stroke=self.config["redaction_color"])
                        redact_annot.update()
            
            # Apply redactions
            page.apply_redactions()
        
        doc.save(output_path, garbage=4, deflate=True)
        doc.close()
        
        self._log_redaction(input_path, output_path, pii_data)
        return output_path
    
    def _redact_word(self, input_path, output_path):
        """Enhanced Word document redaction."""
        text = self.extract_text_from_file(input_path)
        pii_data = self.detect_pii(text)
        
        if not pii_data:
            import shutil
            shutil.copy2(input_path, output_path)
            return output_path
        
        doc = docx.Document(input_path)
        
        # Create replacement mapping
        replacements = {}
        for pii_type, matches in pii_data.items():
            for match in matches:
                replacements[match["text"]] = "█" * len(match["text"])
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for original, replacement in replacements.items():
                if original in paragraph.text:
                    paragraph.text = paragraph.text.replace(original, replacement)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for original, replacement in replacements.items():
                        if original in cell.text:
                            cell.text = cell.text.replace(original, replacement)
        
        doc.save(output_path)
        self._log_redaction(input_path, output_path, pii_data)
        return output_path
    
    def _redact_excel(self, input_path, output_path):
        """Enhanced Excel file redaction."""
        text = self.extract_text_from_file(input_path)
        pii_data = self.detect_pii(text)
        
        if not pii_data:
            import shutil
            shutil.copy2(input_path, output_path)
            return output_path
        
        workbook = load_workbook(input_path)
        black_fill = PatternFill(start_color="000000", end_color="000000", fill_type="solid")
        
        # Create replacement mapping
        replacements = {}
        for pii_type, matches in pii_data.items():
            for match in matches:
                replacements[match["text"]] = "REDACTED"
        
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row.cells:
                    if cell.value:
                        cell_text = str(cell.value)
                        for original, replacement in replacements.items():
                            if original in cell_text:
                                cell.value = replacement
                                cell.fill = black_fill
        
        workbook.save(output_path)
        self._log_redaction(input_path, output_path, pii_data)
        return output_path
    
    def _redact_text(self, input_path, output_path):
        """Enhanced text file redaction with precise positioning."""
        text = self.extract_text_from_file(input_path)
        pii_data = self.detect_pii(text)
        
        if not pii_data:
            import shutil
            shutil.copy2(input_path, output_path)
            return output_path
        
        redacted_text = text
        
        # Sort matches by position (reverse order to maintain positions)
        all_matches = []
        for pii_type, matches in pii_data.items():
            for match in matches:
                all_matches.append(match)
        
        all_matches.sort(key=lambda x: x["start"], reverse=True)
        
        for match in all_matches:
            start = match["start"]
            end = match["end"]
            redacted_text = redacted_text[:start] + "█" * (end - start) + redacted_text[end:]
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(redacted_text)
        
        self._log_redaction(input_path, output_path, pii_data)
        return output_path
    
    def _redact_image(self, input_path, output_path):
        """Enhanced image redaction with OCR-based positioning."""
        text = self.extract_text_from_file(input_path)
        pii_data = self.detect_pii(text)
        
        if not pii_data:
            import shutil
            shutil.copy2(input_path, output_path)
            return output_path
        
        # Load image
        img = Image.open(input_path)
        draw = ImageDraw.Draw(img)
        
        # Get OCR data with bounding boxes
        ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        
        # Create mapping of text to bounding boxes
        text_boxes = {}
        for i, word in enumerate(ocr_data['text']):
            if word.strip():
                x = ocr_data['left'][i]
                y = ocr_data['top'][i]
                w = ocr_data['width'][i]
                h = ocr_data['height'][i]
                text_boxes[word] = (x, y, x + w, y + h)
        
        # Redact identified PII
        for pii_type, matches in pii_data.items():
            for match in matches:
                match_text = match["text"]
                
                # Find matching text in OCR results
                for word, box in text_boxes.items():
                    if match_text in word or word in match_text:
                        draw.rectangle(box, fill="black")
        
        # Save redacted image
        img.save(output_path)
        self._log_redaction(input_path, output_path, pii_data)
        return output_path
    
    def _log_redaction(self, input_path, output_path, pii_data):
        """Enhanced logging of redaction details."""
        log_entry = {
            "timestamp": datetime.now().isoformat(),
            "input_file": input_path,
            "output_file": output_path,
            "pii_detected": {k: len(v) for k, v in pii_data.items()},
            "total_pii_count": sum(len(v) for v in pii_data.values()),
            "pii_details": {k: [{"text": "***REDACTED***", "confidence": m["confidence"]} for m in v] for k, v in pii_data.items()}
        }
        
        self.redaction_log.append(log_entry)
        logger.info(f"Redacted {log_entry['total_pii_count']} PII instances from {input_path}")
    
    def get_redaction_summary(self):
        """Get comprehensive summary of all redactions performed."""
        if not self.redaction_log:
            return {"message": "No redactions performed yet"}
        
        summary = {
            "total_files_processed": len(self.redaction_log),
            "total_pii_redacted": sum(log["total_pii_count"] for log in self.redaction_log),
            "pii_types_found": {},
            "redaction_log": self.redaction_log
        }
        
        # Aggregate PII types
        for log in self.redaction_log:
            for pii_type, count in log["pii_detected"].items():
                if pii_type not in summary["pii_types_found"]:
                    summary["pii_types_found"][pii_type] = 0
                summary["pii_types_found"][pii_type] += count
        
        return summary
    
    def validate_redaction_quality(self, original_path, redacted_path):
        """Validate redaction quality by checking for PII remnants."""
        try:
            redacted_text = self.extract_text_from_file(redacted_path)
            remaining_pii = self.detect_pii(redacted_text)
            
            if remaining_pii:
                logger.warning(f"Potential PII remnants found in redacted file: {remaining_pii}")
                return False
            
            return True
        except Exception as e:
            logger.error(f"Error validating redaction quality: {e}")
            return False


def ensure_proper_extension(file_path):
    """Ensure file has proper extension format."""
    path = Path(file_path)
    if not path.suffix:
        return str(path) + '.txt'
    return str(path)


def main():
    """Main function for CLI and API usage."""
    if len(sys.argv) < 2:
        print("Usage: python advanced_pii_redactor.py <input_file> [output_file]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    if output_file:
        output_file = ensure_proper_extension(output_file)
    
    try:
        # Initialize redactor
        redactor = AdvancedPIIRedactor()
        
        # Perform redaction
        result = redactor.redact_file(input_file, output_file)
        
        # Validate redaction quality
        is_valid = redactor.validate_redaction_quality(input_file, result)
        
        # Print results
        print(f"Redacted file saved to: {result}")
        print(f"Redaction quality validated: {'✓' if is_valid else '✗'}")
        
        # Print summary
        summary = redactor.get_redaction_summary()
        print(f"Total PII instances redacted: {summary['total_pii_redacted']}")
        print(f"PII types found: {list(summary['pii_types_found'].keys())}")
        
        # Return success for API integration
        return {
            "success": True,
            "output_file": result,
            "summary": summary,
            "validation_passed": is_valid
        }
        
    except Exception as e:
        logger.error(f"Error: {e}")
        print(f"Error: {e}")
        return {
            "success": False,
            "error": str(e)
        }


if __name__ == "__main__":
    main()